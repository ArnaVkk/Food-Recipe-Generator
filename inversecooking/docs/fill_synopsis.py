# =============================================================================
# ACADEMIC SYNOPSIS GENERATOR
# Bachelor of Engineering - Computer Science (AIML Specialization)
# Chandigarh University, Gharuan, Mohali
# =============================================================================

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

doc = Document()

# =============================================================================
# DOCUMENT STYLING - Times New Roman
# =============================================================================
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

def add_heading_caps(text, level=1):
    """Add bold, all-caps heading"""
    h = doc.add_heading('', level=level)
    run = h.add_run(text.upper())
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14) if level == 1 else Pt(12)
    return h

def add_subheading(text):
    """Add bold subheading"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_body_text(text):
    """Add justified body text"""
    p = doc.add_paragraph(text)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.27)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_placeholder(description):
    """Add image/diagram placeholder"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"[INSERT SCREENSHOT/DIAGRAM HERE: {description}]")
    run.bold = True
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    doc.add_paragraph()
    return p

# =============================================================================
# TITLE PAGE
# =============================================================================
for _ in range(3):
    doc.add_paragraph()

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("SMART INGREDIENT IDENTIFIER: AN AI-POWERED FOOD ANALYSIS SYSTEM USING DEEP LEARNING")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(16)

doc.add_paragraph()
doc.add_paragraph()

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run("A Project Work Synopsis")
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted in the partial fulfillment for the award of the degree of")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("BACHELOR OF ENGINEERING")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("IN")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("COMPUTER SCIENCE WITH SPECIALIZATION IN")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("ARTIFICIAL INTELLIGENCE AND MACHINE LEARNING")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Submitted by:")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

# Student names
students = [
    ("22BCS10363", "Arnav Kumar"),
    ("22BCS10364", "Ansh Singh"),
    ("22BCS10365", "Dayal Chandra Sati")
]

for uid, name in students:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{uid} — {name}")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("Under the Supervision of:")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("<Project Supervisor Name>")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("CHANDIGARH UNIVERSITY")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("GHARUAN, MOHALI – 140413, PUNJAB, INDIA")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run("January, 2026")
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

doc.add_page_break()

# =============================================================================
# ABSTRACT
# =============================================================================
h = doc.add_heading('', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run("ABSTRACT")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

abstract_text = """The automated recognition and classification of food items from digital images presents a significant challenge in the domain of computer vision due to the inherent variability in food presentation, preparation styles, and cultural diversity. This project proposes the development of a Smart Ingredient Identifier, an artificial intelligence-powered food analysis system that leverages deep learning techniques to accurately identify food dishes from images and subsequently provide detailed ingredient information and cooking recipes.

The proposed system employs a Convolutional Neural Network (CNN) based on the EfficientNet-B0 architecture, utilizing transfer learning from pre-trained ImageNet weights to achieve optimal feature extraction capabilities. The model is trained on a comprehensive dataset comprising 181 food categories, including 80 traditional Indian cuisine dishes and 101 Western/International dishes, totaling 113,900 training images. A two-phase training methodology is implemented, consisting of a warmup phase for classifier initialization followed by a fine-tuning phase with cosine annealing learning rate scheduling.

Experimental evaluation demonstrates that the proposed model achieves a validation accuracy of 84.8% with a minimal training-validation gap of +2.5%, indicating robust generalization capabilities and effective mitigation of overfitting. The system is deployed as an interactive web-based application using the Gradio framework, enabling real-time food recognition and recipe retrieval functionality. The proposed solution establishes a viable foundation for intelligent food analysis systems with potential applications in nutritional management, dietary tracking, allergen detection, and smart kitchen environments."""

p = doc.add_paragraph(abstract_text)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for run in p.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_paragraph()

p = doc.add_paragraph()
run = p.add_run("Keywords: ")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
run2 = p.add_run("Convolutional Neural Networks, Deep Learning, EfficientNet, Transfer Learning, Food Image Classification, Ingredient Recognition, Computer Vision, Gradio Web Application")
run2.font.name = 'Times New Roman'
run2.font.size = Pt(12)

doc.add_page_break()

# =============================================================================
# TABLE OF CONTENTS
# =============================================================================
h = doc.add_heading('', level=1)
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = h.add_run("TABLE OF CONTENTS")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

doc.add_paragraph()

toc_items = [
    ("Title Page", "i"),
    ("Abstract", "ii"),
    ("Table of Contents", "iii"),
    ("1. Introduction", "1"),
    ("    1.1 Problem Definition", "1"),
    ("    1.2 Project Overview", "3"),
    ("    1.3 Hardware Specification", "5"),
    ("    1.4 Software Specification", "6"),
    ("2. Literature Survey", ""),
    ("    2.1 Existing System", ""),
    ("    2.2 Proposed System", ""),
    ("    2.3 Literature Review Summary", ""),
    ("3. Problem Formulation", ""),
    ("4. Objectives", ""),
    ("5. Methodology", ""),
    ("6. Experimental Setup", ""),
    ("7. Conclusion", ""),
    ("8. Tentative Chapter Plan", ""),
    ("References", ""),
]

toc_table = doc.add_table(rows=len(toc_items), cols=2)
for i, (item, page) in enumerate(toc_items):
    cell1 = toc_table.rows[i].cells[0]
    cell2 = toc_table.rows[i].cells[1]
    
    p1 = cell1.paragraphs[0]
    run1 = p1.add_run(item)
    run1.font.name = 'Times New Roman'
    run1.font.size = Pt(12)
    
    p2 = cell2.paragraphs[0]
    p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run2 = p2.add_run(page)
    run2.font.name = 'Times New Roman'
    run2.font.size = Pt(12)

doc.add_page_break()

# =============================================================================
# 1. INTRODUCTION
# =============================================================================
h = doc.add_heading('', level=1)
run = h.add_run("1. INTRODUCTION")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

intro_para1 = """The rapid advancement of artificial intelligence and machine learning technologies has revolutionized numerous domains, with computer vision emerging as one of the most impactful applications in contemporary technological landscapes. Food image analysis, a specialized subset of computer vision, has garnered significant research attention due to its multifaceted applications in healthcare informatics, nutritional management systems, dietary tracking applications, and intelligent lifestyle solutions. The exponential proliferation of food delivery applications, digital recipe platforms, and health-conscious consumer behavior has accentuated the demand for automated systems capable of interpreting and analyzing food images with high accuracy and reliability."""

add_body_text(intro_para1)

intro_para2 = """Traditional approaches to food recognition relied predominantly on handcrafted visual features, including color histograms, texture descriptors, and shape-based representations. While these methodologies demonstrated acceptable performance under controlled laboratory conditions, they exhibited substantial sensitivity to variations in illumination, camera perspective, food presentation, and environmental factors, thereby limiting their practical applicability in real-world scenarios. The paradigm shift towards deep learning architectures, particularly Convolutional Neural Networks (CNNs), has fundamentally transformed the landscape of image classification and object recognition, enabling the development of robust food recognition systems with unprecedented accuracy levels."""

add_body_text(intro_para2)

intro_para3 = """Despite these technological advancements, existing food recognition models are characterized by significant limitations pertaining to cultural bias and cuisine-specific generalization. The majority of publicly available food datasets, including the widely-used Food-101 and Recipe1M+, exhibit a pronounced bias towards Western culinary traditions, resulting in trained models that demonstrate substantially degraded performance when confronted with culturally diverse food items, particularly those originating from South Asian, East Asian, or Middle Eastern cuisines. This inherent bias in training data distribution represents a critical impediment to the global deployment of food recognition technologies."""

add_body_text(intro_para3)

add_placeholder("High-level System Architecture Diagram showing the flow from Image Input → Preprocessing → CNN Model → Prediction → Recipe Database → Output")

# =============================================================================
# 1.1 PROBLEM DEFINITION
# =============================================================================
h = doc.add_heading('', level=2)
run = h.add_run("1.1 PROBLEM DEFINITION")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

prob_para1 = """The fundamental problem addressed in this research endeavor pertains to the inadequate capability of contemporary food recognition models to accurately classify dishes originating from diverse culinary traditions, with particular emphasis on Indian cuisine. Empirical investigation of the existing inversecooking model, developed by Facebook AI Research and trained on the Recipe1M+ dataset, revealed a critical failure in recognizing traditional Indian food items. When presented with an image of Garlic Naan, a staple Indian flatbread, the model erroneously classified it as "SALMON" with a confidence score of 78.3%."""

add_body_text(prob_para1)

add_placeholder("Screenshot showing the misclassification of Garlic Naan as Salmon with 78.3% confidence")

prob_para2 = """This misclassification phenomenon can be attributed to several interconnected factors inherent in the model's training paradigm:"""

p = doc.add_paragraph(prob_para2)
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
for run in p.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

# Bullet points
bullets = [
    "Dataset Bias: The Recipe1M+ dataset predominantly comprises Western and American cuisine items, with minimal representation of South Asian food categories, resulting in a skewed feature space that fails to capture the visual characteristics of Indian dishes.",
    "Feature Mismatch: The absence of Indian food items during the training phase implies that the model's learned feature representations are not optimized for recognizing textures, colors, and shapes characteristic of Indian culinary preparations.",
    "Cultural Homogeneity: The training corpus exhibits a lack of cultural diversity, leading to models that generalize poorly across geographically and culturally distinct food traditions.",
    "Visual Feature Confusion: Without exposure to relevant training examples, the model resorts to matching input images with visually similar but semantically incorrect categories from its training distribution."
]

for bullet in bullets:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(bullet)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_paragraph()

prob_para3 = """The model demonstrates a systematic inability to recognize a comprehensive range of traditional Indian dishes, including but not limited to: Biryani, Butter Chicken, Chicken Tikka Masala, Naan, Dosa, Idli, Samosa, Paneer Tikka, Dal Makhani, Palak Paneer, Chole Bhature, Gulab Jamun, Jalebi, Rasgulla, Ras Malai, and numerous other regional specialties. This deficiency renders such models impractical for deployment in geographical regions characterized by diverse culinary practices and poses a significant barrier to the development of globally applicable food recognition solutions."""

add_body_text(prob_para3)

prob_para4 = """Furthermore, the problem extends beyond mere classification accuracy to encompass the broader challenge of developing machine learning systems that are culturally inclusive and capable of serving diverse user populations without inherent biases that disadvantage specific demographic groups."""

add_body_text(prob_para4)

# =============================================================================
# 1.2 PROJECT OVERVIEW
# =============================================================================
h = doc.add_heading('', level=2)
run = h.add_run("1.2 PROJECT OVERVIEW")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

overview_para1 = """This project proposes the development of a Smart Ingredient Identifier system that addresses the aforementioned limitations through the implementation of a culturally inclusive deep learning model trained on a comprehensive, multi-cuisine dataset. The proposed solution leverages transfer learning techniques with the EfficientNet-B0 architecture, combined with a diverse training corpus encompassing both Indian and Western food categories, to achieve robust cross-cultural food recognition capabilities."""

add_body_text(overview_para1)

overview_para2 = """The project was executed through an iterative development methodology, with multiple model versions developed and evaluated to systematically address challenges related to overfitting and generalization. The evolutionary progression of model development is documented as follows:"""

add_body_text(overview_para2)

add_placeholder("Graph showing Training vs Validation Accuracy across Model Versions V1, V2, V3, and Final Model")

# Model evolution table
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Table 1: Model Evolution and Performance Metrics")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

model_table = doc.add_table(rows=5, cols=5)
model_table.style = 'Table Grid'
headers = ["Model Version", "Dataset Configuration", "Training Accuracy", "Validation Accuracy", "Generalization Gap"]
data = [
    ["Version 1 (V1)", "90 classes, 3,150 images", "96.2%", "65.5%", "+30.7% (Severe Overfitting)"],
    ["Version 2 (V2)", "90 classes, 3,150 images", "34.2%", "45.6%", "-11.4% (Underfitting)"],
    ["Version 3 (V3)", "90 classes, 3,150 images", "79.1%", "63.6%", "+15.5% (Moderate Overfitting)"],
    ["Final Model", "181 classes, 113,900 images", "87.2%", "84.8%", "+2.4% (Optimal)"],
]

for j, header in enumerate(headers):
    cell = model_table.rows[0].cells[j]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for i, row_data in enumerate(data):
    for j, cell_data in enumerate(row_data):
        cell = model_table.rows[i+1].cells[j]
        cell.text = cell_data
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

overview_para3 = """The analysis of model evolution reveals a critical insight: the primary determinant of model generalization is the scale and diversity of the training dataset rather than the intensity of regularization techniques. Version 1 exhibited severe overfitting due to insufficient training data, while Version 2 demonstrated underfitting resulting from excessive regularization. The final model, trained on a substantially larger and more diverse dataset comprising 113,900 images across 181 categories, achieved optimal performance with minimal overfitting."""

add_body_text(overview_para3)

overview_para4 = """The proposed system architecture encompasses the following key components:"""

p = doc.add_paragraph(overview_para4)
for run in p.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

components = [
    "Deep Learning Model: EfficientNet-B0 backbone with custom classification head, utilizing transfer learning from ImageNet pre-trained weights and a two-phase training strategy (warmup followed by fine-tuning).",
    "Training Dataset: A comprehensive corpus comprising the Indian Food Images Dataset (80 categories from Kaggle) and the Food-101 Dataset (101 Western/International categories from ETH Zurich), totaling 181 distinct food categories.",
    "Web-Based Interface: An interactive Gradio-powered web application enabling real-time image upload, food classification, confidence score visualization, and recipe retrieval functionality.",
    "Recipe Database: A structured repository containing 181 complete recipes with ingredient lists and step-by-step cooking instructions for all supported food categories."
]

for comp in components:
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(comp)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)

doc.add_paragraph()

add_placeholder("Complete System Architecture Diagram showing Image Acquisition Module → Preprocessing Module → Deep Learning Module → Recipe Mapping Module → Web Interface Module")

# =============================================================================
# 1.3 HARDWARE SPECIFICATION
# =============================================================================
h = doc.add_heading('', level=2)
run = h.add_run("1.3 HARDWARE SPECIFICATION")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

hw_para1 = """The development and training of deep learning models for image classification necessitates specialized hardware configurations capable of handling computationally intensive operations, particularly matrix multiplications and convolution operations inherent in Convolutional Neural Networks. The hardware specifications employed in this project are detailed in Table 2."""

add_body_text(hw_para1)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Table 2: Hardware Specification for Model Development")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

hw_table = doc.add_table(rows=8, cols=3)
hw_table.style = 'Table Grid'
hw_headers = ["Component", "Minimum Requirement", "Configuration Used"]
hw_data = [
    ["Graphics Processing Unit (GPU)", "NVIDIA GPU with 4GB VRAM, CUDA Compute Capability 6.0+", "NVIDIA GeForce RTX 3050 Laptop GPU (4GB GDDR6 VRAM)"],
    ["CUDA Toolkit", "CUDA 11.0 or higher", "CUDA 12.4"],
    ["cuDNN Library", "cuDNN 8.0 or higher", "cuDNN 8.9.7"],
    ["System Memory (RAM)", "8 GB DDR4", "16 GB DDR4"],
    ["Storage", "50 GB available space (SSD recommended)", "512 GB NVMe SSD"],
    ["Central Processing Unit (CPU)", "Intel Core i5 / AMD Ryzen 5 or equivalent", "Intel Core i7 / AMD Ryzen 7"],
    ["Display", "1920×1080 resolution", "1920×1080 Full HD"],
]

for j, header in enumerate(hw_headers):
    cell = hw_table.rows[0].cells[j]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for i, row_data in enumerate(hw_data):
    for j, cell_data in enumerate(row_data):
        cell = hw_table.rows[i+1].cells[j]
        cell.text = cell_data
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

hw_para2 = """The utilization of GPU acceleration through NVIDIA's CUDA (Compute Unified Device Architecture) platform is essential for achieving practical training times. The RTX 3050 GPU, featuring Ampere architecture with 2048 CUDA cores, enabled the completion of the final model training (25 epochs on 113,900 images) in approximately 8.5 hours, a duration that would extend to several days on CPU-only configurations."""

add_body_text(hw_para2)

hw_para3 = """For inference and deployment purposes, the hardware requirements are substantially reduced, with the trained model capable of executing real-time predictions on systems equipped with modest GPU configurations or even CPU-only environments, albeit with increased latency (2-5 seconds per prediction on CPU versus sub-second inference on GPU)."""

add_body_text(hw_para3)

# =============================================================================
# 1.4 SOFTWARE SPECIFICATION
# =============================================================================
h = doc.add_heading('', level=2)
run = h.add_run("1.4 SOFTWARE SPECIFICATION")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)

sw_para1 = """The software ecosystem employed in this project comprises a carefully curated stack of open-source tools, libraries, and frameworks optimized for deep learning research and deployment. The software specifications are enumerated in Table 3."""

add_body_text(sw_para1)

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run("Table 3: Software Specification and Technology Stack")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(11)
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

sw_table = doc.add_table(rows=12, cols=3)
sw_table.style = 'Table Grid'
sw_headers = ["Category", "Software/Library", "Version/Details"]
sw_data = [
    ["Operating System", "Microsoft Windows", "Windows 10/11 (64-bit)"],
    ["Programming Language", "Python", "3.10.x"],
    ["Deep Learning Framework", "PyTorch", "2.6.0 with CUDA 12.4 support"],
    ["Computer Vision Library", "TorchVision", "0.21.0"],
    ["CNN Architecture", "EfficientNet-B0", "Pre-trained on ImageNet (1000 classes)"],
    ["Web Framework", "Gradio", "4.0+ (Interactive ML Web Applications)"],
    ["Image Processing", "Pillow (PIL Fork)", "10.0+"],
    ["Numerical Computing", "NumPy", "1.24+"],
    ["Data Manipulation", "Pandas", "2.0+"],
    ["Progress Visualization", "tqdm", "4.65+"],
    ["Development Environment", "Visual Studio Code", "Latest with Python Extension"],
]

for j, header in enumerate(sw_headers):
    cell = sw_table.rows[0].cells[j]
    cell.text = header
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
    cell.paragraphs[0].runs[0].font.size = Pt(10)

for i, row_data in enumerate(sw_data):
    for j, cell_data in enumerate(row_data):
        cell = sw_table.rows[i+1].cells[j]
        cell.text = cell_data
        cell.paragraphs[0].runs[0].font.name = 'Times New Roman'
        cell.paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

sw_para2 = """PyTorch was selected as the primary deep learning framework due to its dynamic computational graph construction, intuitive Pythonic API, extensive pre-trained model repository (TorchVision), and robust GPU acceleration capabilities through native CUDA integration. The EfficientNet-B0 architecture, accessed through TorchVision's model zoo, provides an optimal balance between computational efficiency and classification accuracy, making it particularly suitable for deployment in resource-constrained environments."""

add_body_text(sw_para2)

sw_para3 = """The Gradio framework facilitates rapid development of interactive web interfaces for machine learning models, enabling the creation of user-friendly applications without requiring extensive web development expertise. The framework provides built-in components for image upload, text display, and confidence visualization, significantly accelerating the deployment pipeline from trained model to functional web application."""

add_body_text(sw_para3)

add_placeholder("Screenshot of the Gradio Web Application Interface showing image upload area, prediction results, and recipe display")

doc.add_page_break()

# =============================================================================
# REMAINING SECTIONS - EMPTY PLACEHOLDERS
# =============================================================================
h = doc.add_heading('', level=1)
run = h.add_run("2. LITERATURE SURVEY")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

h = doc.add_heading('', level=2)
run = h.add_run("2.1 EXISTING SYSTEM")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
doc.add_paragraph()

h = doc.add_heading('', level=2)
run = h.add_run("2.2 PROPOSED SYSTEM")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
doc.add_paragraph()

h = doc.add_heading('', level=2)
run = h.add_run("2.3 LITERATURE REVIEW SUMMARY")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(12)
doc.add_paragraph()

doc.add_page_break()

h = doc.add_heading('', level=1)
run = h.add_run("3. PROBLEM FORMULATION")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
doc.add_paragraph()

h = doc.add_heading('', level=1)
run = h.add_run("4. OBJECTIVES")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
doc.add_paragraph()

h = doc.add_heading('', level=1)
run = h.add_run("5. METHODOLOGY")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
doc.add_paragraph()

h = doc.add_heading('', level=1)
run = h.add_run("6. EXPERIMENTAL SETUP")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
doc.add_paragraph()

h = doc.add_heading('', level=1)
run = h.add_run("7. CONCLUSION")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)
doc.add_paragraph()

h = doc.add_heading('', level=1)
run = h.add_run("8. TENTATIVE CHAPTER PLAN FOR THE PROPOSED WORK")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

chapters = [
    "CHAPTER 1: INTRODUCTION",
    "CHAPTER 2: LITERATURE REVIEW",
    "CHAPTER 3: OBJECTIVES",
    "CHAPTER 4: SYSTEM DESIGN AND METHODOLOGY",
    "CHAPTER 5: IMPLEMENTATION AND EXPERIMENTAL SETUP",
    "CHAPTER 6: RESULTS AND ANALYSIS",
    "CHAPTER 7: CONCLUSION AND FUTURE SCOPE"
]

for ch in chapters:
    p = doc.add_paragraph(ch)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

doc.add_page_break()

h = doc.add_heading('', level=1)
run = h.add_run("REFERENCES")
run.bold = True
run.font.name = 'Times New Roman'
run.font.size = Pt(14)

refs = [
    "[1]",
    "[2]",
    "[3]",
    "[4]",
    "[5]",
    "[6]",
    "[7]",
]

for ref in refs:
    p = doc.add_paragraph(ref)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)

# =============================================================================
# SAVE DOCUMENT
# =============================================================================
output_path = 'Synopsis_Filled.docx'
doc.save(output_path)

print("=" * 70)
print("ACADEMIC SYNOPSIS GENERATED SUCCESSFULLY")
print("=" * 70)
print(f"\nOutput File: {output_path}")
print(f"University: Chandigarh University, Gharuan, Mohali")
print(f"Degree: B.E. Computer Science (AIML Specialization)")
print("\n" + "=" * 70)
print("FILLED SECTIONS (Academic Formatting Applied):")
print("=" * 70)
print("  ✓ Title Page (Times New Roman, Centered)")
print("  ✓ Abstract (Justified, 12pt)")
print("  ✓ Keywords (AIML Technical Terms)")
print("  ✓ Table of Contents")
print("  ✓ 1. INTRODUCTION (Bold, All Caps Heading)")
print("  ✓ 1.1 PROBLEM DEFINITION (With placeholder for screenshot)")
print("  ✓ 1.2 PROJECT OVERVIEW (With Table 1: Model Evolution)")
print("  ✓ 1.3 HARDWARE SPECIFICATION (Table 2: Hardware Specs)")
print("  ✓ 1.4 SOFTWARE SPECIFICATION (Table 3: Software Stack)")
print("\n" + "=" * 70)
print("PLACEHOLDER TAGS INSERTED FOR:")
print("=" * 70)
print("  • High-level System Architecture Diagram")
print("  • Screenshot of Naan→Salmon Misclassification")
print("  • Training vs Validation Accuracy Graph")
print("  • Complete System Architecture Diagram")
print("  • Gradio Web Application Interface Screenshot")
print("\n" + "=" * 70)
print("EMPTY SECTIONS (For Manual Completion):")
print("=" * 70)
print("  - 2. Literature Survey (2.1, 2.2, 2.3)")
print("  - 3. Problem Formulation")
print("  - 4. Objectives")
print("  - 5. Methodology")
print("  - 6. Experimental Setup")
print("  - 7. Conclusion")
print("  - 8. Tentative Chapter Plan")
print("  - References")
